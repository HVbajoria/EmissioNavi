import streamlit as st
import pandas as pd
import numpy as np
import os
import joblib
import hashlib
import matplotlib.pyplot as plt
import matplotlib
from managed_db import *
import lime
import lime.lime_tabular
from docx import Document

def gradient_text(text, color1, color2):
    gradient_css = f"""
        background: -webkit-linear-gradient(left, {color1}, {color2});
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-weight: bold;
        font-size: 42px;
    """
    return f'<span style="{gradient_css}">{text}</span>'

color1 = "#0d3270"
color2 = "#0fab7b"
text = "Trans Sustain"

styled_text = gradient_text(text, color1, color2)
st.write(f"<div>{styled_text}</div>", unsafe_allow_html=True)

st.write("Calculate and Log your environmental impact while shipping and commuting using Machine learning algorithms: Logistic Regression and Decision Tree Classifier.")

matplotlib.use('Agg')

st.header(':green[Enter the details:]')

feature_names_best = ['ton', 'unit', 'transportation_mode', 'c_carbon_neutral', 'eco_friendly', 'delivery', 'ascites','varices', 'items_count', 'avg_cost', 'distance', 'approx_weight', 'mode_number','b_carbon_neutral']

distance_dict = {"Miles":1,"Kilometers":2}
feature_dict = {"No":1,"Yes":2}
transportmodel_dict = {"Rail":1,"Ship":2, "Air":3, "Truck":4, "Car":5}

def get_value(val,my_dict):
	for key,value in my_dict.items():
		if val == key:
			return value

def get_fvalue(val):
	feature_dict = {"No":1,"Yes":2}
	for key,value in feature_dict.items():
		if val == key:
			return value

def get_tvalue(val):
	for key,value in transportmodel_dict.items():
		if val == key:
			return value

# Load ML Models
def load_model(model_file):
	loaded_model = joblib.load(open(os.path.join(model_file),"rb"))
	return loaded_model

doc = Document()
doc.add_heading('Your Carbon Emission', level=1)

st.write("Average shipment weights:")
st.write(":green[Freight tons per railcar - 90:] ")
st.write(":green[Freight tons per truck - 16:]")

transportation_mode = st.radio("Mode of transportation", tuple(transportmodel_dict.keys()))
ton = st.number_input("Tons you wish to ship", 1, 10000)
unit = st.radio("Unit of measurement for distance", tuple(distance_dict.keys()))
distance = st.number_input("Average distance which will be travelled", 1, 1000)
mode_number = 1

st.write("Additional Information about your Shipment!")
items_count = st.number_input("Number of items chosen? (If buying retail)", 0, 100000)
avg_cost = st.number_input("Average cost of products", 0, 100000)

approx_weight = st.number_input("Average weight of packing material used in packaging", 0, 100000)

eco_friendly = st.radio("Have you selected eco-friendly packaging?", tuple(feature_dict.keys()))

delivery = st.radio("Will it be same-day delivery?", tuple(feature_dict.keys()))
ascites = st.radio("Is the product refurbished or new?", tuple(feature_dict.keys()))
varices = st.radio("Can the product be recyled?", tuple(feature_dict.keys()))
b_carbon_neutral = st.radio("Is the business carbon-neutral?", tuple(feature_dict.keys()))
c_carbon_neutral = st.radio("Is the shipment company carbon-neutral?", tuple(feature_dict.keys()))

feature_list = [ton, get_value(unit, distance_dict), get_tvalue(transportation_mode), get_fvalue(c_carbon_neutral),
                get_fvalue(eco_friendly), get_fvalue(delivery), get_fvalue(ascites), get_fvalue(varices),
                items_count, avg_cost, distance, approx_weight, int(mode_number), get_fvalue(b_carbon_neutral)]

pretty_result = {"age": ton, "sex": unit, "steroid": transportation_mode, "antivirals": c_carbon_neutral,
                                 "fatigue": eco_friendly, "spiders": delivery, "ascites": ascites, "varices": varices,
                                 "bilirubin": items_count, "alk_phosphate": avg_cost, "sgot": distance,
                                 "albumin": approx_weight, "protime": mode_number, "histolog": b_carbon_neutral}

single_sample = np.array(feature_list).reshape(1, -1)

model_choice = st.selectbox("Select Model", ["LR", "DecisionTree"])

if st.button("Predict"):
    if model_choice == "DecisionTree":
        loaded_model = load_model("models/decision_tree.pkl")
        prediction = loaded_model.predict(single_sample)
        pred_prob = loaded_model.predict_proba(single_sample)
    else:
        loaded_model = load_model("models/logistic_regression.pkl")
        prediction = loaded_model.predict(single_sample)
        pred_prob = loaded_model.predict_proba(single_sample)
		
    doc.add_heading('Chosen mode of transporttaion: '+transportation_mode, level=3)
    doc.add_heading('Tons of shipment you wish to ship: '+str(ton), level=3)
    doc.add_heading('Chosen unit of measurement for distance: '+unit, level=3)
    doc.add_heading('Average distance which will be travelled: '+transportation_mode, level=3)
    doc.add_heading('Additional Information about your Shipment: ', level=2)
    doc.add_heading('Number of items chosen? (If buying retail): '+str(items_count), level=3)
    doc.add_heading('Average cost of products: '+str(avg_cost), level=3)
    doc.add_heading('Average weight of packing material used in packaging: '+str(approx_weight), level=3)
    doc.add_heading('Average weight of packing material used in packaging: '+eco_friendly, level=3)
    doc.add_heading('It will be same-day delivery: '+delivery, level=3)
    doc.add_heading('The product refurbished: '+ascites, level=3)
    doc.add_heading('The product can be recyled: '+varices, level=3)
    doc.add_heading('The business is carbon-neutral: '+b_carbon_neutral, level=3)
    doc.add_heading('The shipment company is carbon-neutral: '+c_carbon_neutral, level=3)
    doc.add_heading('Machine Learning Model Choice: '+model_choice, level=3)
    doc.add_heading('Result', level=1)
	
    if unit == "Miles":
        prescriptive_message_temp ="""
	    <div style="overflow-x: auto; padding:10px;border-radius:5px;margin:10px;">
		<h3 style="text-align:justify;color:black;padding:10px">Options to offset your carbon footprint</h3>
		<ul>
		<li style="text-align:justify;color:black;padding:10px">By not selecting same-day delivery you could reduce your CO2e emissions by an estimated 0.8 tons </li>
		<li style="text-align:justify;color:black;padding:10px">Overall Carbon footprint score is equivalent to 1 car off the road for 1 hour</li>
		<li style="text-align:justify;color:black;padding:10px">Plant one tree today!</li>
		<li style="text-align:justify;color:black;padding:10px">Use refurbuished items instead of new</li>
		<ul>
	    </div>
	    """
        st.warning(":red[You have significant carbon footprint!:]")
        doc.add_heading('You have significant carbon footprint!', level=2)
		
        if(transportation_mode == 'Car'):
            st.write(":red[Total estimated CO2e emissions from all selected modes] "+str(pred_prob[0][0]*0.7))
            st.write(":red[Overall Carbon Footprint score] "+str(items_count*1.2))
            doc.add_heading("Total estimated CO2e emissions from all selected modes "+str(pred_prob[0][0]*0.7)+"\nOverall Carbon Footprint score "+str(items_count*1.2), level=3)
            st.subheader("What can you do to offset this carbon footprint?")
            st.markdown(prescriptive_message_temp, unsafe_allow_html=True)
        
        else:
            doc.add_heading("Total estimated CO2e emissions from all selected modes "+str(pred_prob[0][0]*0.2)+"\nOverall Carbon Footprint score "+str(items_count*0.4), level=3)
            st.write(":red[Total estimated CO2e emissions from all selected modes] "+str(pred_prob[0][0]*0.2))
            st.write(":red[Overall Carbon Footprint score] "+str(items_count*0.4))
            st.subheader("What can you do to offset this carbon footprint?")
            st.markdown(prescriptive_message_temp, unsafe_allow_html=True)
            doc.add_heading("Options to offset your carbon footprint",level=2)
            doc.add_heading("1. By not selecting same-day delivery you could reduce your CO2e emissions by an estimated 0.8 tons", level=3)
            doc.add_heading("2. Overall Carbon footprint score is equivalent to 1 car off the road for 1 hour", level=3)
            doc.add_heading("3. Use refurbuished items instead of new", level=3)
            doc.add_heading("4. Plant one tree today!", level=3)
    else:
        st.success("You do not have significant carbon footprint results")
        doc.add_heading("You do not have significant carbon footprint results", level=2)
        doc.add_heading("Performance Score impact "+str(items_count),level=3)
        st.write(":green[Performance Score Impact:] "+str(items_count))
        st.subheader("Prediction Probability Score using {}".format(model_choice))

    doc.save('Report.doc')
    st.download_button(
    label="Download Report",
    data=open("Report.doc", "rb").read(),
    file_name="Report.doc",
    mime="application/octet-stream",
    help="Click to download the carbon emission."
    )
	
footer="""<style>

a:hover,  a:active {
color: red;
background-color: transparent;
text-decoration: underline;
}

.footer {
position: fixed;
left: 0;
bottom: 0;
width: 100%;
background-color: white;
color: black;
text-align: center;
}
</style>
<div class="footer">
<p>Developed with ❤️ by <a style='display: inline; text-align: center;' href="https://www.linkedin.com/in/harshavardhan-bajoria/" target="_blank">Harshavardhan Bajoria</a></p>
</div>
"""
st.markdown(footer,unsafe_allow_html=True)